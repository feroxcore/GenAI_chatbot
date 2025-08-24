import os
import sys
import json
import logging
import sqlite3
import requests
import base64
import re
import socket
import time
import asyncio
import faiss
import pandas as pd
import numpy as np
from fastapi import FastAPI, UploadFile, File, HTTPException, WebSocket, WebSocketDisconnect, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from starlette.requests import Request
from collections import deque
from datetime import datetime
from starlette.websockets import WebSocketState
from pathlib import Path
from typing import List, Optional
import chardet

# PDF processing imports
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None

# --- CONFIG ---
DB_PATH = "chatbot_data.db"
EMBED_DIM = 1536  # text-embedding-3-small dimension
EURON_API_KEY = "euri-52f40fd2263d59037383c38f20e8a4a38e2a0d106fcb0780c3786007c8d6a2b2"

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("chatbot.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# --- VECTOR STORE CLASS ---
class EmbeddingVectorStore:
    def __init__(self, json_dir='chunks', embed_dim=1536):
        self.json_dir = Path(json_dir)
        self.embed_dim = embed_dim
        self.index = faiss.IndexFlatL2(embed_dim)
        self.chunk_metadata = []
        self.is_loaded = False
        
    def add_chunk(self, text: str, embedding: List[float], metadata: dict):
        """Add a single chunk with its embedding to the store"""
        try:
            vector = np.array(embedding).reshape(1, -1).astype('float32')
            self.index.add(vector)
            self.chunk_metadata.append({
                'text': text,
                'metadata': metadata,
                'index_id': self.index.ntotal - 1
            })
        except Exception as e:
            logger.error(f"Error adding chunk to vector store: {e}")
    
    def load_from_json_files(self):
        """Load all embeddings from JSON chunk files"""
        self.index.reset()
        self.chunk_metadata.clear()
        
        if not self.json_dir.exists():
            logger.warning(f"Chunks directory {self.json_dir} does not exist")
            return
        
        embeddings_loaded = 0
        for json_file in self.json_dir.glob('*_chunks.json'):
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                for chunk in data.get('chunks', []):
                    embedding = chunk.get('embedding')
                    if embedding:
                        self.add_chunk(
                            text=chunk['content'],
                            embedding=embedding,
                            metadata=chunk.get('metadata', {})
                        )
                        embeddings_loaded += 1
                        
            except Exception as e:
                logger.error(f"Error loading chunks from {json_file}: {e}")
        
        self.is_loaded = True
        logger.info(f"Loaded {embeddings_loaded} embeddings from {len(list(self.json_dir.glob('*_chunks.json')))} files")
    
    def search(self, query_embedding: List[float], top_k: int = 5) -> List[str]:
        """Search for similar chunks using vector similarity"""
        if not self.is_loaded or self.index.ntotal == 0:
            logger.warning("Vector store is empty or not loaded")
            return []
        
        try:
            query_vector = np.array(query_embedding).reshape(1, -1).astype('float32')
            distances, indices = self.index.search(query_vector, min(top_k, self.index.ntotal))
            
            results = []
            for idx in indices[0]:
                if idx >= 0 and idx < len(self.chunk_metadata):
                    results.append(self.chunk_metadata[idx]['text'])
            
            return results
        except Exception as e:
            logger.error(f"Error during vector search: {e}")
            return []

# Global vector store instance
vector_store = EmbeddingVectorStore(embed_dim=1536)

# --- DB SETUP ---
def init_sqlite_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS feedback (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id TEXT,
        user_query TEXT,
        response TEXT,
        feedback_type TEXT,
        file_name TEXT,
        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')
    conn.commit()
    conn.close()

init_sqlite_db()

def get_db_connection():
    return sqlite3.connect(DB_PATH)

# --- Euron API helpers ---
import time

def get_euron_embeddings(texts):
    """
    Generate embeddings for a list of texts using Euron API
    Returns a list of embeddings (one per input text)
    """
    try:
        url = "https://api.euron.one/api/v1/euri/embeddings"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {EURON_API_KEY}"
        }
        payload = {
            "input": texts,  # Send as list for batch processing
            "model": "text-embedding-3-small"
        }
        
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        response.raise_for_status()  # Raise exception for bad status codes
        
        data = response.json()
        
        # Check if response has the expected structure
        if "data" not in data:
            logger.error("Invalid response structure: missing 'data' field")
            return []
        
        embeddings = []
        for item in data['data']:
            if 'embedding' in item and item['embedding']:
                embeddings.append(item['embedding'])
            else:
                logger.warning("Missing or empty embedding in response item")
                embeddings.append(None)
        
        # Validate we got the expected number of embeddings
        if len(embeddings) != len(texts):
            logger.warning(f"Expected {len(texts)} embeddings, got {len(embeddings)}")
        
        return embeddings
        
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error getting Euron embeddings: {e}")
        return []
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in embedding response: {e}")
        return []
    except Exception as e:
        logger.error(f"Error getting Euron embeddings: {e}")
        return []

def euron_chat_completion(messages):
    """
    Generate chat completion using Euron API
    Accepts OpenAI-style messages format and returns the response content
    """
    try:
        url = "https://api.euron.one/api/v1/euri/chat/completions"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {EURON_API_KEY}"
        }
        
        # Format messages properly - messages should already be in the correct format
        payload = {
            "messages": messages,  # Use messages as-is (should be list of dicts with role/content)
            "model": "gpt-4.1-nano",
            "max_tokens": 1000,
            "temperature": 0.7
        }
        
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        response.raise_for_status()
        
        data = response.json()
        
        # Check response structure and extract content
        if "choices" in data and len(data["choices"]) > 0:
            choice = data["choices"][0]
            if "message" in choice and "content" in choice["message"]:
                return choice["message"]["content"]
            else:
                logger.error("Invalid response structure: missing message content")
                return "I apologize, but I received an invalid response. Please try again."
        else:
            logger.error("Invalid response structure: missing choices")
            return "I apologize, but I didn't receive a proper response. Please try again."
            
    except requests.exceptions.RequestException as e:
        logger.error(f"Request error in Euron chat completion: {e}")
        return "I'm experiencing connection issues. Please try again in a moment."
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error in chat completion: {e}")
        return "I received an invalid response format. Please try again."
    except Exception as e:
        logger.error(f"Error in Euron chat completion: {e}")
        return "I encountered an unexpected error. Please try again."


# --- TEXT PROCESSING ---
def simple_chunk_text(text, chunk_size=300, overlap=80):
    """Simple text chunking by character count with overlap"""
    chunks = []
    words = text.split()
    
    current_chunk = []
    current_size = 0
    
    for word in words:
        current_chunk.append(word)
        current_size += len(word) + 1
        
        if current_size >= chunk_size:
            chunk_text = ' '.join(current_chunk)
            chunks.append(chunk_text)
            
            overlap_words = int(len(current_chunk) * (overlap / chunk_size))
            current_chunk = current_chunk[-overlap_words:] if overlap_words > 0 else []
            current_size = sum(len(word) + 1 for word in current_chunk)
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def save_chunks_with_embeddings(chunks, metadata, file_name, embeddings):
    """Save chunks, metadata, and embeddings to JSON file"""
    try:
        os.makedirs("chunks", exist_ok=True)
        
        base_name = os.path.splitext(file_name)[0]
        json_filename = f"chunks/{base_name}_chunks.json"
        
        chunk_data = {
            "file_name": file_name,
            "processed_at": datetime.now().isoformat(),
            "total_chunks": len(chunks),
            "embedding_model": "text-embedding-3-small",
            "embedding_dim": 1536,
            "chunks": []
        }
        
        for i, (chunk, meta, embedding) in enumerate(zip(chunks, metadata, embeddings)):
            chunk_entry = {
                "chunk_id": i,
                "content": chunk,
                "metadata": meta,
                "embedding": embedding,
                "character_count": len(chunk),
                "word_count": len(chunk.split())
            }
            chunk_data["chunks"].append(chunk_entry)
        
        with open(json_filename, 'w', encoding='utf-8') as f:
            json.dump(chunk_data, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved {len(chunks)} chunks with embeddings to {json_filename}")
        return json_filename
        
    except Exception as e:
        logger.error(f"Error saving chunks with embeddings: {e}")
        return None

# --- FILE PROCESSING ---
def process_file(file_path, file_name, start_chunk_id=0):
    """Complete file processing function supporting multiple formats"""
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    chunks = []
    metadata = []
    chunk_id = start_chunk_id
    
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File {file_path} not found")

        # Excel files
        if ext in [".xlsx", ".xls"]:
            sheets = pd.read_excel(file_path, sheet_name=None)
            df = pd.concat(sheets.values(), ignore_index=True)
            df.fillna('', inplace=True)
            headers = df.columns.tolist()
            header_text = " | ".join(f"{h}" for h in headers)
            
            for index, row in df.iterrows():
                row_content = header_text + "\n" + " | ".join(f"{v}" for v in row.values)
                chunks.append(row_content)
                metadata.append({
                    "type": "table_with_header",
                    "file_name": file_name,
                    "chunk_id": chunk_id,
                    "row_number": index + 1
                })
                chunk_id += 1

        # CSV files
        elif ext == ".csv":
            with open(file_path, "rb") as f:
                raw_data = f.read(100000)
                encoding = chardet.detect(raw_data)["encoding"] or "utf-8"
            
            df = pd.read_csv(file_path, encoding=encoding)
            df.fillna('', inplace=True)
            headers = df.columns.tolist()
            header_text = " | ".join(f"{h}" for h in headers)
            
            for index, row in df.iterrows():
                row_content = header_text + "\n" + " | ".join(f"{v}" for v in row.values)
                chunks.append(row_content)
                metadata.append({
                    "type": "table_with_header",
                    "file_name": file_name,
                    "chunk_id": chunk_id,
                    "row_number": index + 1
                })
                chunk_id += 1

        # PDF files
        elif ext == ".pdf":
            pdf_text = ""
            
            # Method 1: pdfplumber
            if pdfplumber:
                try:
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text() or ""
                            tables = page.extract_tables()
                            table_text = ""
                            for table in tables:
                                for row in table:
                                    row_text = " | ".join(cell if cell is not None else "" for cell in row)
                                    table_text += row_text + "\n"
                            pdf_text += page_text + "\n" + table_text + "\n\n"
                except Exception as e:
                    logger.warning(f"pdfplumber failed for {file_name}: {e}")
            
            # Method 2: PyMuPDF
            if not pdf_text.strip() and fitz:
                try:
                    doc = fitz.open(file_path)
                    pdf_text = "\n".join([page.get_text() for page in doc])
                    doc.close()
                except Exception as e:
                    logger.warning(f"PyMuPDF failed for {file_name}: {e}")
            
            # Method 3: PyPDF2
            if not pdf_text.strip() and PdfReader:
                try:
                    with open(file_path, "rb") as f:
                        reader = PdfReader(f)
                        pdf_text = " ".join([p.extract_text() or "" for p in reader.pages])
                except Exception as e:
                    logger.warning(f"PyPDF2 failed for {file_name}: {e}")
            
            if pdf_text.strip():
                text_chunks = simple_chunk_text(pdf_text)
                for chunk in text_chunks:
                    if chunk.strip():
                        chunks.append(chunk)
                        metadata.append({
                            "type": "text",
                            "file_name": file_name,
                            "chunk_id": chunk_id,
                            "source": "pdf_extraction"
                        })
                        chunk_id += 1

        # DOCX files
        elif ext == ".docx":
            if not docx:
                raise ImportError("python-docx library required for DOCX files")
            
            doc = docx.Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                paragraphs.append(table_text)
            
            full_text = "\n\n".join(paragraphs)
            text_chunks = simple_chunk_text(full_text)
            
            for chunk in text_chunks:
                if chunk.strip():
                    chunks.append(chunk)
                    metadata.append({
                        "type": "text",
                        "file_name": file_name,
                        "chunk_id": chunk_id,
                        "source": "docx_extraction"
                    })
                    chunk_id += 1

        # Text files
        elif ext == ".txt":
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = ""
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text = f.read().strip()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text:
                text_chunks = simple_chunk_text(text)
                for chunk in text_chunks:
                    if chunk.strip():
                        chunks.append(chunk)
                        metadata.append({
                            "type": "text",
                            "file_name": file_name,
                            "chunk_id": chunk_id,
                            "source": "text_file"
                        })
                        chunk_id += 1

        # JSON files
        elif ext == ".json":
            with open(file_path, "r", encoding="utf-8") as f:
                json_data = json.load(f)
            
            json_text = json.dumps(json_data, indent=2, ensure_ascii=False)
            text_chunks = simple_chunk_text(json_text)
            
            for chunk in text_chunks:
                if chunk.strip():
                    chunks.append(chunk)
                    metadata.append({
                        "type": "json_data",
                        "file_name": file_name,
                        "chunk_id": chunk_id,
                        "source": "json_file"
                    })
                    chunk_id += 1

        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {ext}")

        logger.info(f"Successfully processed {file_name}: {len(chunks)} chunks created")
        return chunks, metadata

    except Exception as e:
        logger.error(f"Error processing {file_name}: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing {file_name}: {str(e)}")

def process_file_with_embeddings(file_path, file_name, start_chunk_id=0):
    """Process file and generate embeddings for each chunk"""
    chunks, metadata = process_file(file_path, file_name, start_chunk_id)
    
    if not chunks:
        logger.warning(f"No chunks extracted from {file_name}")
        return chunks, metadata, None
    
    logger.info(f"Generating embeddings for {len(chunks)} chunks...")
    all_embeddings = []
    
    # Process in smaller batches to respect API limits
    batch_size = 20
    total_batches = (len(chunks) - 1) // batch_size + 1
    
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i + batch_size]
        batch_num = i // batch_size + 1
        
        logger.info(f"Processing embedding batch {batch_num}/{total_batches} ({len(batch_chunks)} chunks)")
        
        try:
            # Get embeddings for this batch
            batch_embeddings = get_euron_embeddings(batch_chunks)
            
            if batch_embeddings and len(batch_embeddings) == len(batch_chunks):
                all_embeddings.extend(batch_embeddings)
                logger.info(f"Batch {batch_num} successful: {len(batch_embeddings)} embeddings")
            else:
                logger.error(f"Batch {batch_num} failed: expected {len(batch_chunks)}, got {len(batch_embeddings) if batch_embeddings else 0}")
                # Add None for failed embeddings to maintain alignment
                all_embeddings.extend([None] * len(batch_chunks))
            
            # Rate limiting - small delay between batches
            if i + batch_size < len(chunks):
                time.sleep(0.5)
                
        except Exception as e:
            logger.error(f"Exception in batch {batch_num}: {e}")
            all_embeddings.extend([None] * len(batch_chunks))
    
    # Filter out chunks with failed embeddings
    valid_chunks = []
    valid_metadata = []
    valid_embeddings = []
    
    for chunk, meta, embedding in zip(chunks, metadata, all_embeddings):
        if embedding is not None:
            valid_chunks.append(chunk)
            valid_metadata.append(meta)
            valid_embeddings.append(embedding)
    
    success_rate = len(valid_embeddings) / len(chunks) * 100
    logger.info(f"Embedding success rate: {success_rate:.1f}% ({len(valid_embeddings)}/{len(chunks)})")
    
    if not valid_embeddings:
        logger.error("No valid embeddings generated for any chunks")
        return [], [], None
    
    # Save to JSON
    json_file = save_chunks_with_embeddings(valid_chunks, valid_metadata, file_name, valid_embeddings)
    
    # Add to vector store
    for chunk, embedding, meta in zip(valid_chunks, valid_embeddings, valid_metadata):
        vector_store.add_chunk(chunk, embedding, meta)
    
    logger.info(f"Successfully processed {file_name}: {len(valid_chunks)} chunks with embeddings saved")
    return valid_chunks, valid_metadata, json_file


def search_similar_chunks_from_json(query_text, k=5):
    """Search for similar chunks using embeddings"""
    try:
        query_embeddings = get_euron_embeddings([query_text])
        if not query_embeddings or len(query_embeddings) == 0:
            logger.warning("Could not generate embedding for query")
            return []
        
        if not vector_store.is_loaded:
            vector_store.load_from_json_files()
        
        similar_chunks = vector_store.search(query_embeddings[0], top_k=k)
        return similar_chunks
        
    except Exception as e:
        logger.error(f"Search error: {e}")
        return []

# --- EXCEPTION HANDLING ---
def print_exception_details(e):
    exc_type, exc_obj, exc_tb = sys.exc_info()
    exception_type = exc_type.__name__
    line_no = exc_tb.tb_lineno
    error_message = str(e)
    function_name = exc_tb.tb_frame.f_code.co_name
    file_name = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
    error_final = f"Error_Message: {error_message} | Function: {function_name} | Exception_Type: {exception_type} | File_Name: {file_name} | Line_no.: {line_no}"
    logger.error(error_final)
    return error_final

# --- FASTAPI APP ---
app = FastAPI()
templates = Jinja2Templates(directory="templates")

def decode_base64(data):
    try:
        return base64.b64decode(data).decode('utf-8')
    except Exception as e:
        logger.error("Base64 decode error: %s", str(e))
        return "UNKNOWN"

@app.get("/", response_class=HTMLResponse)
async def get_index(request: Request):
    try:
        return templates.TemplateResponse("index.html", {"request": request})
    except Exception as e:
        error_message = print_exception_details(e)
        raise HTTPException(status_code=500, detail=f"Failed to load index page: {error_message}")

active_connections = set()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

user_sessions = {}

def add_feedback(user_id, user_query, response, feedback_type, file_name=""):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        response = response[:2000] if len(response) > 2000 else response
        
        cursor.execute("""
            INSERT INTO feedback (user_id, user_query, response, feedback_type, file_name, timestamp) 
            VALUES (?, ?, ?, ?, ?, ?)
        """, (user_id, user_query, response, feedback_type, file_name, datetime.now()))
        
        conn.commit()
        logger.info("Feedback added for user %s: type=%s", user_id, feedback_type)
        return True
    except Exception as e:
        error_message = print_exception_details(e)
        logger.error("Error adding feedback for user %s: %s", user_id, error_message)
        return False
    finally:
        if 'conn' in locals():
            conn.close()

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_path = f"uploads/{file.filename}"
        os.makedirs("uploads", exist_ok=True)
        
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        logger.info(f"Processing file: {file.filename}")
        
        chunks, metadata, json_file_path = process_file_with_embeddings(file_path, file.filename)
        
        if not chunks:
            return {
                "status": "error",
                "message": "No content could be extracted from the file"
            }
        
        return {
            "status": "success",
            "filename": file.filename,
            "chunks_processed": len(chunks),
            "chunks_file": json_file_path,
            "embeddings_generated": True,
            "message": f"File processed successfully. {len(chunks)} chunks with embeddings saved."
        }
        
    except Exception as e:
        error_message = print_exception_details(e)
        logger.error(f"Upload error: {error_message}")
        raise HTTPException(status_code=500, detail=error_message)

@app.websocket("/chat")
async def websocket_chat(websocket: WebSocket):
    await websocket.accept()
    active_connections.add(websocket)
    logger.info("Active connections: %d", len(active_connections))
    user_id = str(id(websocket))

    try:
        client_ip_raw = websocket.client.host
        client_ip = client_ip_raw.split(':')[0].strip()
        try:
            raw_host = socket.gethostbyaddr(client_ip)[0]
        except socket.herror:
            raw_host = "Unknown"
        host = raw_host.split('.')[0].upper()
        logger.info("IP = %s, HOST = %s", client_ip, host)
    except Exception as e:
        error_message = print_exception_details(e)
        logger.error("Hostname resolution failed: %s", error_message)
        host = "UNKNOWN"
        client_ip = "UNKNOWN"

    user_sessions[user_id] = {
        "ip": client_ip,
        "host": host,
        "history": deque(maxlen=10)
    }

    try:
        await websocket.send_json({
            "type": "system",
            "message": "Welcome! I'm your AI assistant. How can I help you today?"
        })

        while True:
            try:
                data = await websocket.receive_json()
                logger.info("Received message: %s", data)
                
                message_type = data.get('type', 'message')
                
                if message_type == 'message':
                    user_message = data.get('message', '').strip()
                    if not user_message:
                        continue
                    
                    user_sessions[user_id]["history"].append({"role": "user", "content": user_message})
                    
                    await websocket.send_json({"type": "typing", "status": True})
                    
                    # Search for relevant context using embeddings
                    context_chunks = search_similar_chunks_from_json(user_message, k=3)
                    context = "\n\n".join(context_chunks) if context_chunks else ""
                    
                    system_message = "You are a helpful AI assistant."
                    if context:
                        system_message += f" Use the following context from uploaded documents to answer the user's question:\n\n{context}"
                    
                    messages = [
                        {"role": "system", "content": system_message},
                        *list(user_sessions[user_id]["history"])
                    ]
                    
                    ai_response = euron_chat_completion(messages)
                    
                    user_sessions[user_id]["history"].append({"role": "assistant", "content": ai_response})
                    
                    await websocket.send_json({
                        "type": "message",
                        "message": ai_response,
                        "timestamp": datetime.now().isoformat(),
                        "context_used": len(context_chunks) > 0
                    })
                    
                elif message_type == 'feedback':
                    feedback_type = data.get('feedback_type')
                    
                    if user_sessions[user_id]["history"]:
                        last_response = user_sessions[user_id]["history"][-1]["content"]
                        last_query = user_sessions[user_id]["history"][-2]["content"] if len(user_sessions[user_id]["history"]) > 1 else ""
                        
                        add_feedback(user_id, last_query, last_response, feedback_type)
                        
                        await websocket.send_json({
                            "type": "feedback_received",
                            "message": "Thank you for your feedback!"
                        })
                        
            except WebSocketDisconnect:
                logger.info("Client disconnected normally")
                break
            except Exception as e:
                error_message = print_exception_details(e)
                logger.error("Error processing WebSocket message: %s", error_message)
                await websocket.send_json({
                    "type": "error", 
                    "message": "I encountered an error processing your request. Please try again."
                })
                
    except Exception as e:
        error_message = print_exception_details(e)
        logger.error("WebSocket connection error: %s", error_message)
    finally:
        active_connections.discard(websocket)
        if user_id in user_sessions:
            del user_sessions[user_id]
        logger.info("Active connections after cleanup: %d", len(active_connections))

@app.on_event("startup")
async def startup_event():
    """Load existing embeddings on startup"""
    logger.info("Loading existing embeddings...")
    vector_store.load_from_json_files()
    logger.info("Chatbot ready!")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)